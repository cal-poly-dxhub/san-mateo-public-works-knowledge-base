import io
import pypdf
import docx
import openpyxl


def extract_text(file_bytes, filename):
    """Extract text from various document formats"""
    ext = filename.lower().split('.')[-1]
    
    if ext == 'pdf':
        return extract_pdf(file_bytes)
    elif ext in ['doc', 'docx']:
        return extract_docx(file_bytes)
    elif ext in ['xls', 'xlsx']:
        return extract_xlsx(file_bytes)
    else:
        return file_bytes.decode('utf-8', errors='ignore')


def extract_pdf(file_bytes):
    """Extract text from PDF"""
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        text = []
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text.append(extracted)
        return '\n'.join(text) if text else ""
    except Exception as e:
        print(f"Error extracting PDF: {e}")
        return ""


def extract_docx(file_bytes):
    """Extract text from DOCX"""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    text.append(row_text)
        return '\n'.join(text) if text else ""
    except Exception as e:
        print(f"Error extracting DOCX: {e}")
        return ""


def extract_xlsx(file_bytes):
    """Extract text from XLSX"""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        text = []
        for sheet in wb.worksheets:
            text.append(f"Sheet: {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                row_text = '\t'.join(str(c) if c else '' for c in row)
                if row_text.strip():
                    text.append(row_text)
        return '\n'.join(text) if text else ""
    except Exception as e:
        print(f"Error extracting XLSX: {e}")
        return ""
