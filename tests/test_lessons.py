#!/usr/bin/env python3
"""Lessons Learned Tests"""

import io
import json
import time
import pytest
import requests
from .test_config import API_URL, MAX_PROCESSING_TIMEOUT
from .conftest import get_cached_auth_token


def get_auth_only():
    return {"Authorization": get_cached_auth_token()}


def get_auth_headers():
    return {
        "Authorization": get_cached_auth_token(),
        "Content-Type": "application/json"
    }


def wait_for_lessons(project_name, min_count=1, timeout=30, interval=2):
    """Poll until lessons exist or timeout"""
    start = time.time()
    while time.time() - start < timeout:
        resp = requests.get(f"{API_URL}/projects/{project_name}/lessons-learned", headers=get_auth_only())
        if resp.status_code == 200:
            lessons = resp.json().get("lessons", [])
            if len(lessons) >= min_count:
                return lessons
        time.sleep(interval)
    return []


def test_extract_lessons_from_document():
    """Test extracting lessons from uploaded document"""
    project_name = f"test-lessons-{int(time.time())}"
    
    # Create project
    requests.post(
        f"{API_URL}/setup-wizard",
        headers=get_auth_headers(),
        json={
            "projectName": project_name,
            "projectType": "Other",
            "location": "Test",
            "areaSize": "1.0",
            "specialConditions": []
        }
    )
    
    # Upload document with lessons
    response = requests.post(
        f"{API_URL}/upload-url",
        headers=get_auth_headers(),
        json={
            "files": [{
                "fileName": "lessons.txt",
                "projectName": project_name,
                "projectType": "Other",
                "extractLessons": True
            }]
        }
    )
    
    assert response.status_code == 200
    upload_url = response.json()["uploads"][0]["uploadUrl"]
    
    # Upload content
    content = """
    Lesson 1: Early utility coordination prevents project delays.
    Lesson 2: Budget contingency of 15% is essential.
    Lesson 3: Weekly stakeholder meetings improve communication.
    """
    requests.put(upload_url, data=content.encode())
    
    # Wait for processing with polling
    lessons = wait_for_lessons(project_name, min_count=3)
    assert len(lessons) >= 3
    
    # Cleanup
    requests.delete(f"{API_URL}/projects/{project_name}", headers=get_auth_only())


def test_get_project_lesson_conflicts():
    """Test getting lesson conflicts for a project"""
    project_name = f"test-conflicts-{int(time.time())}"
    
    # Create project
    requests.post(
        f"{API_URL}/setup-wizard",
        headers=get_auth_headers(),
        json={
            "projectName": project_name,
            "projectType": "Other",
            "location": "Test",
            "areaSize": "1.0",
            "specialConditions": []
        }
    )
    
    # Upload conflicting documents
    for i in range(2):
        response = requests.post(
            f"{API_URL}/upload-url",
            headers=get_auth_headers(),
            json={
                "files": [{
                    "fileName": f"conflict_{i}.txt",
                    "projectName": project_name,
                    "projectType": "Other",
                    "extractLessons": True
                }]
            }
        )
        
        if response.status_code == 200:
            upload_url = response.json()["uploads"][0]["uploadUrl"]
            if i == 0:
                content = "Lesson: Early coordination is critical for success."
            else:
                content = "Lesson: Late coordination can be more cost-effective."
            requests.put(upload_url, data=content.encode())
    
    wait_for_lessons(project_name, min_count=2)
    
    # Get conflicts
    response = requests.get(
        f"{API_URL}/projects/{project_name}/conflicts",
        headers=get_auth_only()
    )
    
    assert response.status_code == 200
    result = response.json()
    assert "conflicts" in result or isinstance(result, list)
    
    # Cleanup
    requests.delete(f"{API_URL}/projects/{project_name}", headers=get_auth_only())


def test_resolve_project_lesson_conflict():
    """Test resolving a lesson conflict"""
    project_name = f"test-resolve-{int(time.time())}"
    
    # Create project with conflicting lessons
    requests.post(
        f"{API_URL}/setup-wizard",
        headers=get_auth_headers(),
        json={
            "projectName": project_name,
            "projectType": "Other",
            "location": "Test",
            "areaSize": "1.0",
            "specialConditions": []
        }
    )
    
    # Upload conflicting documents
    for i in range(2):
        response = requests.post(
            f"{API_URL}/upload-url",
            headers=get_auth_headers(),
            json={
                "files": [{
                    "fileName": f"conflict_{i}.txt",
                    "projectName": project_name,
                    "projectType": "Other",
                    "extractLessons": True
                }]
            }
        )
        
        if response.status_code == 200:
            upload_url = response.json()["uploads"][0]["uploadUrl"]
            content = f"Lesson {i}: Conflicting information about project management."
            requests.put(upload_url, data=content.encode())
    
    wait_for_lessons(project_name, min_count=2)
    
    # Get conflicts
    conflicts_response = requests.get(
        f"{API_URL}/projects/{project_name}/conflicts",
        headers=get_auth_only()
    )
    
    if conflicts_response.status_code == 200:
        conflicts = conflicts_response.json().get("conflicts", [])
        
        if len(conflicts) > 0:
            conflict_id = conflicts[0].get("id") or conflicts[0].get("conflict_id")
            
            # Resolve conflict
            response = requests.post(
                f"{API_URL}/projects/{project_name}/conflicts/resolve",
                headers=get_auth_headers(),
                json={
                    "conflict_id": conflict_id,
                    "resolution": "keep_existing"
                }
            )
            
            assert response.status_code == 200
    
    # Cleanup
    requests.delete(f"{API_URL}/projects/{project_name}", headers=get_auth_only())


def test_extract_lessons_from_pdf():
    """Test extracting lessons from a PDF document"""
    from reportlab.pdfgen import canvas
    
    project_name = f"test-pdf-{int(time.time())}"
    
    requests.post(
        f"{API_URL}/setup-wizard",
        headers=get_auth_headers(),
        json={"projectName": project_name, "projectType": "Other", "location": "Test", "areaSize": "1.0", "specialConditions": []}
    )
    
    response = requests.post(
        f"{API_URL}/upload-url",
        headers=get_auth_headers(),
        json={"files": [{"fileName": "lessons.pdf", "projectName": project_name, "projectType": "Other", "extractLessons": True}]}
    )
    assert response.status_code == 200
    upload_url = response.json()["uploads"][0]["uploadUrl"]
    
    # Create PDF with actual text content
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer)
    c.drawString(100, 700, "Lesson learned: PDF coordination with utilities prevents delays.")
    c.drawString(100, 680, "Lesson learned: Budget 20% contingency for unexpected conditions.")
    c.save()
    pdf_buffer.seek(0)
    requests.put(upload_url, data=pdf_buffer.getvalue())
    
    lessons = wait_for_lessons(project_name, min_count=1)
    assert len(lessons) >= 1
    
    requests.delete(f"{API_URL}/projects/{project_name}", headers=get_auth_only())


def test_extract_lessons_from_docx():
    """Test extracting lessons from a DOCX document"""
    from docx import Document
    
    project_name = f"test-docx-{int(time.time())}"
    
    requests.post(
        f"{API_URL}/setup-wizard",
        headers=get_auth_headers(),
        json={"projectName": project_name, "projectType": "Other", "location": "Test", "areaSize": "1.0", "specialConditions": []}
    )
    
    response = requests.post(
        f"{API_URL}/upload-url",
        headers=get_auth_headers(),
        json={"files": [{"fileName": "lessons.docx", "projectName": project_name, "projectType": "Other", "extractLessons": True}]}
    )
    assert response.status_code == 200
    upload_url = response.json()["uploads"][0]["uploadUrl"]
    
    # Create DOCX with lessons
    doc = Document()
    doc.add_paragraph("Lesson 1: DOCX early coordination prevents project delays.")
    doc.add_paragraph("Lesson 2: DOCX budget contingency of 15% is essential.")
    doc.add_paragraph("Lesson 3: DOCX weekly stakeholder meetings improve communication.")
    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    requests.put(upload_url, data=docx_buffer.getvalue())
    
    lessons = wait_for_lessons(project_name, min_count=3)
    assert len(lessons) >= 3
    
    requests.delete(f"{API_URL}/projects/{project_name}", headers=get_auth_only())


def test_extract_lessons_from_xlsx():
    """Test extracting lessons from an XLSX spreadsheet"""
    from openpyxl import Workbook
    
    project_name = f"test-xlsx-{int(time.time())}"
    
    requests.post(
        f"{API_URL}/setup-wizard",
        headers=get_auth_headers(),
        json={"projectName": project_name, "projectType": "Other", "location": "Test", "areaSize": "1.0", "specialConditions": []}
    )
    
    response = requests.post(
        f"{API_URL}/upload-url",
        headers=get_auth_headers(),
        json={"files": [{"fileName": "lessons.xlsx", "projectName": project_name, "projectType": "Other", "extractLessons": True}]}
    )
    assert response.status_code == 200
    upload_url = response.json()["uploads"][0]["uploadUrl"]
    
    # Create XLSX with lessons
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "Lesson"
    ws["A2"] = "Lesson 1: XLSX early coordination prevents project delays."
    ws["A3"] = "Lesson 2: XLSX budget contingency of 15% is essential."
    ws["A4"] = "Lesson 3: XLSX weekly stakeholder meetings improve communication."
    
    xlsx_buffer = io.BytesIO()
    wb.save(xlsx_buffer)
    requests.put(upload_url, data=xlsx_buffer.getvalue())
    
    lessons = wait_for_lessons(project_name, min_count=3)
    assert len(lessons) >= 3
    
    requests.delete(f"{API_URL}/projects/{project_name}", headers=get_auth_only())


def test_get_master_lesson_project_types():
    """Test getting available project types for master lessons"""
    response = requests.get(
        f"{API_URL}/lessons/project-types",
        headers=get_auth_only()
    )
    
    assert response.status_code == 200
    result = response.json()
    assert isinstance(result, (list, dict))


def test_get_master_lessons_by_type():
    """Test getting aggregated lessons by project type"""
    response = requests.get(
        f"{API_URL}/lessons/by-type/Other",
        headers=get_auth_only()
    )
    
    assert response.status_code == 200
    result = response.json()
    assert "lessons" in result or isinstance(result, list)


def test_update_master_lesson():
    """Test updating a master lesson"""
    # Get lessons first
    response = requests.get(
        f"{API_URL}/lessons/by-type/Other",
        headers=get_auth_only()
    )
    
    if response.status_code == 200:
        lessons = response.json().get("lessons", [])
        
        if len(lessons) > 0:
            lesson_id = lessons[0].get("id") or lessons[0].get("lesson_id")
            
            # Update lesson
            update_response = requests.put(
                f"{API_URL}/lessons/{lesson_id}",
                headers=get_auth_headers(),
                json={
                    "title": "Updated Master Lesson",
                    "lesson": "Updated content",
                    "impact": "Updated impact",
                    "recommendation": "Updated recommendation"
                }
            )
            
            assert update_response.status_code in [200, 404]


def test_get_master_lesson_conflicts_by_type():
    """Test getting master lesson conflicts by type"""
    response = requests.get(
        f"{API_URL}/lessons/conflicts/by-type/Other",
        headers=get_auth_only()
    )
    
    assert response.status_code == 200
    result = response.json()
    assert "conflicts" in result or isinstance(result, list)


def test_resolve_master_lesson_conflict():
    """Test resolving a master lesson conflict"""
    # Get conflicts first
    response = requests.get(
        f"{API_URL}/lessons/conflicts/by-type/Other",
        headers=get_auth_only()
    )
    
    if response.status_code == 200:
        conflicts = response.json().get("conflicts", [])
        
        if len(conflicts) > 0:
            conflict_id = conflicts[0].get("id") or conflicts[0].get("conflict_id")
            
            # Resolve conflict - API expects 'decision' and 'project_type'
            resolve_response = requests.post(
                f"{API_URL}/lessons/conflicts/resolve/{conflict_id}",
                headers=get_auth_headers(),
                json={"decision": "keep_existing", "project_type": "Other"}
            )
            
            assert resolve_response.status_code in [200, 404]
